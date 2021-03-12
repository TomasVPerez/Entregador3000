#TomasPerez
#https://github.com/TomasVPerez

import os, time
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

class Sistema:
    def __init__(self):
        self.extensiones = ["jpg", "PNG", "jpeg"]
        self.imagenes = []

        self.usuario = usuario = os.path.expanduser("~")
        self.descargas = descargas = f"{usuario}/Downloads"

    def mostrarMenu(self):
        print('''
    ______      __                            __              _____ ____  ____  ____ 
   / ____/___  / /_________  ____ _____ _____/ /___  _____   |__  // __ \/ __ \/ __ \\
  / __/ / __ \/ __/ ___/ _ \/ __ `/ __ `/ __  / __ \/ ___/    /_ </ / / / / / / / / /
 / /___/ / / / /_/ /  /  __/ /_/ / /_/ / /_/ / /_/ / /      ___/ / /_/ / /_/ / /_/ / 
/_____/_/ /_/\__/_/   \___/\__, /\__,_/\__,_/\____/_/      /____/\____/\____/\____/  
                          /____/         
by: Tomas Perez
https://github.com/TomasVPerez
''')
        nombreEntrega = input("Nombre del entregable: ")
        header = input("Materia-NombreyApellido(nroAlumno)-Fecha: ")
        self.crearDocumento(nombreEntrega, header)
        
    def guardarImagenes(self):
        archivos = next(os.walk(self.descargas))[2] #Recorre todos los archivos del directorio especificado
        for archivo in archivos:
            if "WhatsApp" in archivo and archivo.split(".")[-1] in self.extensiones: #Si encuentra una foto con 'WhatsApp' en el nombre y las extensiones aceptadas, lo guarda en la lista imagenes.
                self.imagenes.append(archivo)

    def crearDocumento(self, nombre, header):
        word = Document()
        encabezado = word.sections[0].header
        encabezado.paragraphs[0].text = header
        estilo = word.styles["Normal"]
        fuente = estilo.font
        fuente.name = "Arial"
        fuente.size = Pt(11)
        for imagen in self.imagenes:
            creacion = time.ctime(os.path.getctime(f"{self.descargas}/{imagen}"))
            if time.ctime()[:13] == creacion[:13]: #Compara dia mes y hora para corrobrar que son imagenes de wpp descargadas recientemente (sino pondria todas las imagenes de wpp que encuentre en el descarga).
                word.add_picture(f"{self.descargas}/{imagen}", width=Inches(5.0)) #Detecta el orden de descarga automaticamente y agrega las fotos en orden.
                ultimaFoto = word.paragraphs[-1] 
                ultimaFoto.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                pass
        nuevaCarpeta = f"{self.usuario}/Desktop/{nombre}"
        os.mkdir(nuevaCarpeta)
        word.save(f"{nuevaCarpeta}/{nombre}.docx")
        print("Documento creado en el escritorio")
        time.sleep(1)       


def main():
    s = Sistema()
    s.guardarImagenes()
    s.mostrarMenu()

if __name__ == '__main__':
    main()



